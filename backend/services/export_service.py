"""
MoM - Export Service
Generate dokumen MoM dalam 4 format: PDF, DOCX, TXT, MD
Masing-masing dengan desain/template tersendiri.
"""

import os
import io
import logging
import tempfile
from datetime import datetime
from typing import Optional

logger = logging.getLogger(__name__)


def _safe_date(date_str: str) -> str:
    """Format ISO date string menjadi tampilan yang bagus."""
    try:
        dt = datetime.fromisoformat(date_str.replace("Z", "+00:00"))
        return dt.strftime("%d %B %Y, %H:%M WIB")
    except Exception:
        return date_str or datetime.now().strftime("%d %B %Y")


def _clean_text(text: str) -> str:
    """Bersihkan karakter non-ASCII yang tidak didukung font standard PDF."""
    if not text:
        return ""
    # Map common problematic characters to ASCII equivalents
    replacements = {
        "\u2014": "---",  # em-dash
        "\u2013": "--",   # en-dash
        "\u201c": '"',    # smart quote open
        "\u201d": '"',    # smart quote close
        "\u2018": "'",    # smart single quote
        "\u2019": "'",    # smart single quote
        "\u2022": "*",    # bullet point
        "\u2026": "...",  # ellipsis
        "\u2010": "-",    # hyphen
        "\u2011": "-",    # non-breaking hyphen
    }
    # Explicitly clean everything to ASCII (0-127) for FPDF standard fonts
    return "".join(c if ord(c) < 128 else replacements.get(c, ("-" if ord(c) in [8211, 8212] else "?")) for c in text)


class ExportService:
    """Service untuk export MoM ke berbagai format dokumen."""

    # ─────────────────────────────────────────────────────────────
    # PDF Export — Format Formal/Korporat
    # ─────────────────────────────────────────────────────────────
    def export_pdf(self, meeting: dict) -> bytes:
        """
        Generate PDF profesional dengan cover page, section numbering,
        dan tabel action items berwarna.
        """
        try:
            from fpdf import FPDF, XPos, YPos

            class MoMPDF(FPDF):
                def __init__(self, title=""):
                    super().__init__()
                    self.mom_title = title
                    self.set_auto_page_break(auto=True, margin=25)

                def header(self):
                    if self.page_no() > 1:
                        self.set_font("Helvetica", "B", 8)
                        self.set_text_color(150, 150, 150)
                        self.cell(0, 8, f"Minutes of Meeting --- {self.mom_title}", align="L")
                        self.set_text_color(0, 0, 0)
                        self.ln(1)
                        self.set_draw_color(220, 220, 220)
                        self.line(10, self.get_y(), 200, self.get_y())
                        self.ln(4)

                def footer(self):
                    if self.page_no() > 0:
                        self.set_y(-15)
                        self.set_font("Helvetica", "I", 8)
                        self.set_text_color(150, 150, 150)
                        self.cell(0, 10, f"Halaman {self.page_no()} | Dicetak: {datetime.now().strftime('%d/%m/%Y %H:%M')}", align="C")

            title_raw = meeting.get("title", "Minutes of Meeting")
            title = _clean_text(title_raw)
            pdf = MoMPDF(title=title)
            pdf.set_left_margin(20)
            pdf.set_right_margin(20)

            # ── CONTENT ───────────────────────────────────────
            pdf.add_page()
            
            # Header info (Compact)
            pdf.set_font("Helvetica", "B", 18)
            pdf.set_text_color(15, 23, 42)
            pdf.multi_cell(0, 10, title, align="L")
            
            pdf.ln(2)
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(100, 100, 100)
            
            date_str = _safe_date(meeting.get("date", ""))
            lang = "Bahasa Indonesia" if meeting.get("language") == "id" else "English"
            meta_text = f"Tanggal: {date_str} | Bahasa: {lang}"
            
            duration_sec = meeting.get("duration_seconds", 0)
            if duration_sec:
                h = int(duration_sec // 3600)
                m = int((duration_sec % 3600) // 60)
                d_str = f"{h}j {m}m" if h > 0 else f"{m}m"
                meta_text += f" | Durasi: {d_str}"
            
            pdf.cell(0, 7, _clean_text(meta_text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            participants = meeting.get("participants", [])
            if participants:
                pdf.set_font("Helvetica", "I", 9)
                pdf.multi_cell(0, 6, _clean_text("Peserta: " + ", ".join(participants)))
            
            pdf.ln(4)
            pdf.set_draw_color(220, 220, 220)
            pdf.line(20, pdf.get_y(), 190, pdf.get_y())
            pdf.ln(6)

            # MoM Body
            mom_doc = meeting.get("mom_document", "")
            if mom_doc:
                self._render_mom_to_pdf(pdf, _clean_text(mom_doc))

            # Action Items
            action_items = meeting.get("action_items", [])
            if action_items:
                pdf.ln(10)
                # Check space
                if pdf.get_y() > 220:
                    pdf.add_page()
                self._render_action_items_pdf(pdf, action_items)

            return bytes(pdf.output())

        except ImportError:
            logger.error("[Export PDF] fpdf2 not installed")
            raise RuntimeError("fpdf2 library tidak tersedia. Install: pip install fpdf2")
        except Exception as e:
            logger.error(f"[Export PDF] Error: {e}")
            raise

    def _render_mom_to_pdf(self, pdf, mom_text: str):
        """Render teks MoM ke PDF dengan formatting dasar."""
        from fpdf import XPos, YPos

        lines = mom_text.split("\n")
        section_num = 0

        for line in lines:
            line = line.strip()
            if not line:
                pdf.ln(3)
                continue

            # H1 heading (# Title)
            if line.startswith("# "):
                section_num += 1
                pdf.set_font("Helvetica", "B", 15)
                pdf.set_text_color(15, 23, 42)
                pdf.set_fill_color(241, 245, 249)
                pdf.cell(0, 10, f"{section_num}. {line[2:]}", fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(2)

            # H2 heading (## SubTitle)
            elif line.startswith("## "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(51, 65, 85)
                pdf.cell(0, 8, line[3:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                pdf.ln(1)

            # H3 heading
            elif line.startswith("### "):
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(71, 85, 105)
                pdf.cell(0, 7, line[4:], new_x=XPos.LMARGIN, new_y=YPos.NEXT)

            # Bullet points (handles - , * , • )
            elif line.startswith(("- ", "* ", "• ")):
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(79, 70, 229) # Indigo bullet
                pdf.cell(8, 6, chr(149), align="C") # Bullet character
                
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(30, 41, 59)
                pdf.multi_cell(0, 6, line[2:].strip(), markdown=True)
                pdf.ln(2)

            # Numbered list
            elif len(line) > 2 and line[0].isdigit() and (line[1] == "." or (len(line) > 2 and line[2] == ".")):
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(30, 41, 59)
                dot_pos = line.find(".")
                pdf.cell(8, 6, line[:dot_pos+1], align="L")
                
                pdf.set_font("Helvetica", "", 10)
                pdf.multi_cell(0, 6, line[dot_pos+1:].strip(), markdown=True)
                pdf.ln(2)

            # Normal paragraph
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(30, 41, 59)
                pdf.multi_cell(0, 6, line, markdown=True)
                pdf.ln(2)

    def _render_action_items_pdf(self, pdf, action_items: list):
        """Render action items sebagai tabel berwarna di halaman terpisah."""
        from fpdf import XPos, YPos

        # Section header
        pdf.set_fill_color(15, 23, 42)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", "B", 13)
        pdf.cell(0, 12, "  ACTION ITEMS", fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(4)

        # Table header (A4 available width is ~170mm with 20mm margins)
        headers = ["No", "Task", "PIC", "Deadline", "Priority", "Status"]
        col_widths = [8, 82, 20, 20, 20, 20] # Total 170mm
        priority_colors = {"high": (254, 226, 226), "medium": (254, 243, 199), "low": (220, 252, 231)}

        pdf.set_font("Helvetica", "B", 9)
        pdf.set_fill_color(71, 85, 105)
        pdf.set_text_color(255, 255, 255)
        for i, (h, w) in enumerate(zip(headers, col_widths)):
            last = (i == len(headers) - 1)
            pdf.cell(w, 8, h, border=0, fill=True, new_x=XPos.LMARGIN if last else XPos.RIGHT, new_y=YPos.NEXT if last else YPos.TOP)
        pdf.ln(1)

        # Rows
        for idx, item in enumerate(action_items):
            prio = item.get("priority", "medium")
            bg = priority_colors.get(prio, (248, 250, 252))
            pdf.set_fill_color(*bg)
            pdf.set_text_color(30, 41, 59)
            pdf.set_font("Helvetica", "", 8) # Smaller font for rows to fit more text

            status_icon = "Done" if item.get("status") == "done" else "Pending"
            
            task_text = _clean_text(item.get("task", ""))
            # Limit slightly but much more generous than before
            if len(task_text) > 90:
                task_text = task_text[:87] + "..."

            values = [
                str(idx + 1),
                task_text,
                _clean_text(item.get("pic", "-"))[:15],
                _clean_text(item.get("deadline", "-")).replace("TBD", "-")[:15],
                prio.capitalize(),
                status_icon
            ]

            for i, (v, w) in enumerate(zip(values, col_widths)):
                last = (i == len(values) - 1)
                pdf.cell(w, 7, v, border="B", fill=True, new_x=XPos.LMARGIN if last else XPos.RIGHT, new_y=YPos.NEXT if last else YPos.TOP)

        pdf.ln(5)
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(100, 100, 100)
        pdf.cell(0, 5, _clean_text(f"Total: {len(action_items)} action items | {sum(1 for i in action_items if i.get('status')=='done')} selesai"))

    # ─────────────────────────────────────────────────────────────
    # DOCX Export — Format Editable Office
    # ─────────────────────────────────────────────────────────────
    def export_docx(self, meeting: dict) -> bytes:
        """
        Generate DOCX dengan styles profesional, table of contents,
        dan tabel action items yang bisa diedit di Word/Google Docs.
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Cm, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement

            doc = Document()

            # ── Page margins ──────────────────────────────────────
            for section in doc.sections:
                section.top_margin = Cm(2.5)
                section.bottom_margin = Cm(2.5)
                section.left_margin = Cm(2.8)
                section.right_margin = Cm(2.8)

            # ── Cover / Header block ──────────────────────────────
            title = meeting.get("title", "Minutes of Meeting")
            heading = doc.add_heading("MINUTES OF MEETING", level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = heading.runs[0]
            run.font.color.rgb = RGBColor(15, 23, 42)

            # Meeting title
            title_para = doc.add_heading(title, level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata table
            meta_table = doc.add_table(rows=4, cols=2)
            meta_table.style = "Table Grid"
            meta_data = [
                ("Tanggal", _safe_date(meeting.get("date", ""))),
                ("Peserta", ", ".join(meeting.get("participants", []))),
                ("Durasi", f"{meeting.get('duration_seconds', 0) // 60} menit"),
                ("Bahasa", "Bahasa Indonesia" if meeting.get("language") == "id" else "English"),
            ]
            for row, (key, val) in zip(meta_table.rows, meta_data):
                row.cells[0].text = key
                row.cells[1].text = val
                row.cells[0].paragraphs[0].runs[0].bold = True

            doc.add_paragraph()

            # ── MoM Document Content ──────────────────────────────
            mom_doc = meeting.get("mom_document", "")
            if mom_doc:
                self._render_mom_to_docx(doc, mom_doc)

            # ── Action Items Table ────────────────────────────────
            action_items = meeting.get("action_items", [])
            if action_items:
                doc.add_page_break()
                doc.add_heading("Action Items", level=1)

                table = doc.add_table(rows=1, cols=6)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                headers = ["No", "Task", "PIC", "Deadline", "Priority", "Status"]
                hdr_row = table.rows[0]
                for i, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
                    cell.text = hdr
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                for idx, item in enumerate(action_items):
                    row = table.add_row()
                    vals = [
                        str(idx + 1),
                        item.get("task", ""),
                        item.get("pic", "-"),
                        item.get("deadline", "-"),
                        item.get("priority", "medium").capitalize(),
                        "✓ Done" if item.get("status") == "done" else "○ Pending",
                    ]
                    for cell, val in zip(row.cells, vals):
                        cell.text = val

            # ── Footer note ───────────────────────────────────────
            doc.add_paragraph()
            footer_para = doc.add_paragraph(f"Dokumen ini dihasilkan otomatis oleh MoM AI Assistant — {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            footer_para.runs[0].font.size = Pt(8)
            footer_para.runs[0].font.color.rgb = RGBColor(150, 150, 150)
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Save to bytes
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()

        except ImportError:
            logger.error("[Export DOCX] python-docx not installed")
            raise RuntimeError("python-docx tidak tersedia. Install: pip install python-docx")

    def _render_mom_to_docx(self, doc, mom_text: str):
        """Render teks MoM ke paragraf DOCX dengan heading levels."""
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        for line in mom_text.split("\n"):
            line = line.strip()
            if not line:
                doc.add_paragraph()
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith(("- ", "* ")):
                para = doc.add_paragraph(style="List Bullet")
                self._add_inline_formatted_text(para, line[2:])
            else:
                para = doc.add_paragraph()
                self._add_inline_formatted_text(para, line)

    def _add_inline_formatted_text(self, paragraph, text: str):
        """Helper untuk menambahkan text dengan inline bold formatting ke docx paragraph."""
        import re
        # Regex to find **bold** parts
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    # ─────────────────────────────────────────────────────────────
    # TXT Export — Format Plain/Arsip
    # ─────────────────────────────────────────────────────────────
    def export_txt(self, meeting: dict) -> bytes:
        """
        Generate plain text dengan ASCII formatting.
        Cocok untuk email, arsip, atau version control.
        """
        lines = []
        width = 72
        sep = "=" * width
        sub = "-" * width

        title = meeting.get("title", "Minutes of Meeting")
        date_str = _safe_date(meeting.get("date", ""))
        participants = meeting.get("participants", [])
        duration = meeting.get("duration_seconds", 0)
        language = "Bahasa Indonesia" if meeting.get("language") == "id" else "English"

        # Duration Logic
        duration_sec = meeting.get("duration_seconds", 0)
        h = int(duration_sec // 3600)
        m = int((duration_sec % 3600) // 60)
        d_str = f"{h} jam {m} menit" if h > 0 else f"{m} menit"

        # Header
        lines += [
            sep,
            "MINUTES OF MEETING".center(width),
            title.center(width),
            sep,
            "",
            f"Tanggal   : {date_str}",
            f"Durasi    : {d_str}",
            f"Bahasa    : {language}",
        ]
        if participants:
            lines.append(f"Peserta   : {', '.join(participants)}")

        lines += ["", sub, ""]

        # MoM Document
        mom_doc = meeting.get("mom_document", "")
        if mom_doc:
            # Clean markdown symbols for plain text
            for line in mom_doc.split("\n"):
                clean = line.strip()
                if clean.startswith("# "):
                    lines += ["", sub, f"  {clean[2:].upper()}", sub]
                elif clean.startswith("## "):
                    lines += ["", f"  {clean[3:]}", "  " + "-" * (len(clean[3:]) + 2)]
                elif clean.startswith("### "):
                    lines += [f"    {clean[4:]}"]
                elif clean.startswith(("- ", "* ")):
                    lines.append(f"    • {clean[2:]}")
                else:
                    clean_md = clean.replace("**", "").replace("*", "").replace("`", "")
                    lines.append(f"  {clean_md}")

        # Action Items
        action_items = meeting.get("action_items", [])
        if action_items:
            lines += ["", sep, "  ACTION ITEMS", sep, ""]
            for idx, item in enumerate(action_items, 1):
                status = "[✓]" if item.get("status") == "done" else "[ ]"
                lines.append(f"  {idx:2}. {status} {item.get('task', '')}")
                if item.get("pic"):
                    lines.append(f"       PIC      : {item['pic']}")
                if item.get("deadline"):
                    lines.append(f"       Deadline : {item['deadline']}")
                lines.append(f"       Prioritas: {item.get('priority', 'medium').capitalize()}")
                lines.append("")

        # Footer
        lines += [
            sub,
            f"  Dibuat oleh MoM AI Assistant — {datetime.now().strftime('%d/%m/%Y %H:%M')}",
            sub,
        ]

        return "\n".join(lines).encode("utf-8")

    # ─────────────────────────────────────────────────────────────
    # MD Export — Format Developer/GitHub/Notion
    # ─────────────────────────────────────────────────────────────
    def export_md(self, meeting: dict) -> bytes:
        """
        Generate Markdown dengan GFM extensions.
        Action items sebagai checkboxes. Render bagus di GitHub, Notion, Obsidian.
        """
        title = meeting.get("title", "Minutes of Meeting")
        date_str = _safe_date(meeting.get("date", ""))
        participants = meeting.get("participants", [])
        duration = meeting.get("duration_seconds", 0)
        language = "🇮🇩 Bahasa Indonesia" if meeting.get("language") == "id" else "🇺🇸 English"

        duration_sec = meeting.get("duration_seconds", 0)
        h = int(duration_sec // 3600)
        m = int((duration_sec % 3600) // 60)
        d_str = f"{h} jam {m} menit" if h > 0 else f"{m} menit"

        lines = [
            f"# 📋 {title}",
            "",
            "> **Minutes of Meeting** — Dihasilkan otomatis oleh YOTA AI",
            "",
            "## 📌 Informasi Meeting",
            "",
            "| Field | Detail |",
            "|---|---|",
            f"| 📅 Tanggal | {date_str} |",
            f"| ⏱️ Durasi | {d_str} |",
            f"| 🌐 Bahasa | {language} |",
        ]

        if participants:
            lines.append(f"| 👥 Peserta | {', '.join(participants)} |")

        lines += ["", "---", ""]

        # MoM Content — pass through as-is (already Markdown)
        mom_doc = meeting.get("mom_document", "")
        if mom_doc:
            lines.append(mom_doc.strip())
            lines += ["", "---", ""]

        # Action Items as GFM checkboxes
        action_items = meeting.get("action_items", [])
        if action_items:
            lines += ["## ✅ Action Items", ""]

            # Group by priority
            for priority in ["high", "medium", "low"]:
                priority_items = [i for i in action_items if i.get("priority") == priority]
                if priority_items:
                    emoji = {"high": "🔴", "medium": "🟡", "low": "🟢"}.get(priority, "⚪")
                    lines.append(f"### {emoji} {priority.capitalize()} Priority")
                    lines.append("")
                    for item in priority_items:
                        checked = "x" if item.get("status") == "done" else " "
                        task = item.get("task", "")
                        pic = f"**{item['pic']}**" if item.get("pic") else "_TBD_"
                        deadline = item.get("deadline", "_-_")
                        lines.append(f"- [{checked}] {task}")
                        lines.append(f"  - 👤 PIC: {pic} | 🗓️ Deadline: {deadline}")
                    lines.append("")

        # Footer metadata
        lines += [
            "---",
            "",
            f"> *Dokumen ini dibuat otomatis oleh **MoM AI Assistant** pada {datetime.now().strftime('%d %B %Y pukul %H:%M')}.*",
            f"> *Meeting ID: `{meeting.get('id', 'unknown')}`*",
        ]

        return "\n".join(lines).encode("utf-8")

    # ─────────────────────────────────────────────────────────────
    # Unified Export Entry Point
    # ─────────────────────────────────────────────────────────────
    def export(self, meeting: dict, format: str) -> tuple[bytes, str, str]:
        """
        Export meeting ke format tertentu.
        Returns: (content_bytes, media_type, filename)
        """
        # Clean segments for filename to avoid Latin-1 header errors
        raw_title = meeting.get("title", "meeting")
        clean_title = _clean_text(raw_title).replace(" ", "_")
        title_slug = "".join(c for c in clean_title if c.isalnum() or c == "_")[:30]
        
        date_slug = datetime.now().strftime("%Y%m%d")
        base_name = f"MoM_{date_slug}_{title_slug}"

        if format == "pdf":
            return self.export_pdf(meeting), "application/pdf", f"{base_name}.pdf"
        elif format == "docx":
            return self.export_docx(meeting), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{base_name}.docx"
        elif format == "txt":
            return self.export_txt(meeting), "text/plain; charset=utf-8", f"{base_name}.txt"
        elif format == "md":
            return self.export_md(meeting), "text/markdown; charset=utf-8", f"{base_name}.md"
        elif format == "raw":
            raw_text = meeting.get("raw_transcript", "Tidak ada transkrip mentah.")
            return raw_text.encode("utf-8"), "text/plain; charset=utf-8", f"{base_name}_RAW.txt"
        else:
            raise ValueError(f"Format tidak didukung: {format}. Gunakan: pdf, docx, txt, md, raw")


# Singleton
export_service = ExportService()
